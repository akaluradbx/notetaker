import json
import requests
import os
from dotenv import load_dotenv, dotenv_values
import sys

# loading variables from .env file
load_dotenv()

transcription = ['']

from docx import Document

databricks_token=os.getenv("DATABRICKS_TOKEN")
databricks_url=os.getenv("DATABRICKS_URL")


headers = {'Authorization': f'Bearer {databricks_token}', 'Content-Type': 'application/json'}

def read_transacibe_file(trans_file):
    transcription = ""
    with open(trans_file,"r") as f:
        transcription = f.read()
        f.close()
    return transcription


def abstract_summary_extraction(transcription):
    print("..abstract_summary_extraction")
    data = {"messages":[
    {
        "role":"system",
        "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
    },
    {
        "role":"user",
        "content": transcription
    }
    ],"temperature":0.5,"top_p":0.95,"max_tokens":4096}

    data_json = json.dumps(data, allow_nan=True)

    response = requests.request(method='POST', headers=headers, url=databricks_url, data=data_json)
    if response.status_code != 200:
        raise Exception(f'Request failed with status {response.status_code}, {response.text}')

    return response.json()['choices'][0]['message']['content']

def key_points_extraction(transcription):
    print("..key_points_extraction")
    data = {"messages":[
    {
        "role":"system",
        "content": "You are a databricks AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
    },
    {
        "role":"user",
        "content": transcription
    }
    ],"temperature":0.5,"top_p":0.95,"max_tokens":4096}

    data_json = json.dumps(data, allow_nan=True)

    response = requests.request(method='POST', headers=headers, url=databricks_url, data=data_json)
    if response.status_code != 200:
        raise Exception(f'Request failed with status {response.status_code}, {response.text}')

    return response.json()['choices'][0]['message']['content']

def action_item_extraction(transcription):
    print("..action_item_extraction")
    data = {"messages":[
    {
        "role":"system",
        "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
    },
    {
        "role":"user",
        "content": transcription
    }
    ],"temperature":0.5,"top_p":0.95,"max_tokens":4096}

    data_json = json.dumps(data, allow_nan=True)

    response = requests.request(method='POST', headers=headers, url=databricks_url, data=data_json)
    if response.status_code != 200:
        raise Exception(f'Request failed with status {response.status_code}, {response.text}')

    return response.json()['choices'][0]['message']['content']

def sentiment_analysis(transcription):
    print("..sentiment_analysis")
    data = {"messages":[
    {
        "role":"system",
        "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
    },
    {
        "role":"user",
        "content": transcription
    }
    ],"temperature":0.5,"top_p":0.95,"max_tokens":4096}

    data_json = json.dumps(data, allow_nan=True)

    response = requests.request(method='POST', headers=headers, url=databricks_url, data=data_json)
    if response.status_code != 200:
        raise Exception(f'Request failed with status {response.status_code}, {response.text}')

    return response.json()['choices'][0]['message']['content']

def followup_email(transcription):
    print("..followup_email")
    data = {"messages":[
    {
        "role":"system",
        "content": "Write an followup  email summarizing text"
    },
    {
        "role":"user",
        "content": transcription
    }
    ],"temperature":0.5,"top_p":0.95,"max_tokens":4096}

    data_json = json.dumps(data, allow_nan=True)

    response = requests.request(method='POST', headers=headers, url=databricks_url, data=data_json)
    if response.status_code != 200:
        raise Exception(f'Request failed with status {response.status_code}, {response.text}')

    return response.json()['choices'][0]['message']['content']

def blog_content(transcription):
    print("..blog_content")
    data = {"messages":[
    {
        "role":"system",
        "content": "You are a professional blogger. I would like you to read the following text and summarize and create a blog content, provide a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
    },
    {
        "role":"user",
        "content": transcription
    }
    ],"temperature":0.5,"top_p":0.95,"max_tokens":4096}

    data_json = json.dumps(data, allow_nan=True)

    response = requests.request(method='POST', headers=headers, url=databricks_url, data=data_json)
    if response.status_code != 200:
        raise Exception(f'Request failed with status {response.status_code}, {response.text}')

    return response.json()['choices'][0]['message']['content']


def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    followup = followup_email(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment,
        'email':followup
    }

def save_as_docx(data, filename):
    doc = Document()
    for key, value in data.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)

def runme(filename):
    audio_file_path = os.getenv("NOTES_DIR") + filename
    #del transcription = transcribe_audio(audio_file_path)
    transcription = read_transacibe_file(audio_file_path)
    data = meeting_minutes(transcription)
    #print(data)
    output_filename=os.path.basename(audio_file_path).replace("txt","docx")
    save_as_docx(data, os.getenv("SUMMARY_DOC_DIR") +output_filename)
    print("output file = " + output_filename)

if __name__ == "__main__":
    #runme("20240409-211540.txt")
    runme(sys.argv[1:][0])
	
